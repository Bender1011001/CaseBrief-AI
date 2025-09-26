from fastapi import FastAPI, UploadFile, File, Depends, Header, HTTPException, status
from fastapi.responses import StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
import os
import tempfile
import uuid
import logging
import json
import re
from dotenv import load_dotenv
from pydantic_settings import BaseSettings
import firebase_admin
from firebase_admin import credentials, auth
from google.cloud import firestore, storage, vision
import vertexai
from vertexai.generative_models import GenerativeModel, GenerationConfig
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
from tenacity import retry, stop_after_attempt, wait_exponential

load_dotenv()

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class Settings(BaseSettings):
    project_id: str = os.getenv("PROJECT_ID")

    model_config = {"env_file": ".env"}

settings = Settings()

# Initialize clients
cred = credentials.ApplicationDefault()
firebase_admin.initialize_app(cred, {"projectId": settings.project_id})
db = firestore.Client()
bucket_name = "casebrief-ai-uploads"
storage_client = storage.Client()
vision_client = vision.ImageAnnotatorClient()
vertexai.init(project=settings.project_id, location="us-central1")
model = GenerativeModel("gemini-1.5-pro")

# Auth dependency
async def get_current_user(authorization: str = Header(None)):
    if not authorization:
        raise HTTPException(401, "Missing token")
    try:
        decoded_token = auth.verify_id_token(authorization.replace("Bearer ", ""))
        return decoded_token["uid"]
    except:
        raise HTTPException(401, "Invalid token")

# Prompts
PROMPT1 = "Extract the key facts from this legal document text: {full_text}. Output as a structured list of bullet points under 'Facts:'."
PROMPT2 = "Based on the facts: {facts} and full text: {full_text}, provide a legal analysis including issues, holdings, and reasoning. Structure as sections: Issues, Holding, Reasoning."
PROMPT3 = "Synthesize the facts: {facts}, analysis: {analysis}, and full text: {full_text} into a complete case brief. Structure with headings: Facts, Procedural History (infer if needed), Issues, Holding, Reasoning, Conclusion."

@retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=4, max=10))
async def generate_content(prompt):
    response = model.generate_content(prompt)
    return response.text

app = FastAPI(title="CaseBrief AI Backend")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "https://casebrief-ai-prod.web.app"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.post("/v1/process/document")
async def process_document(file: UploadFile = File(...), user_id: str = Depends(get_current_user)):
    logger.info(f"Processing document for user {user_id}")
    if not file.filename.endswith('.pdf'):
        raise HTTPException(400, "Only PDF files allowed")
    if file.size > 10 * 1024 * 1024:
        raise HTTPException(400, "File too large (>10MB)")
    doc_id = str(uuid.uuid4())
    doc_ref = db.collection("users").document(user_id).collection("documents").document(doc_id)
    doc_ref.set({"filename": file.filename, "userId": user_id, "status": "processing"})
    content = await file.read()
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
        temp_file.write(content)
        temp_path = temp_file.name
    # Extract text
    full_text = ""
    with fitz.open(temp_path) as pdf_doc:
        for page in pdf_doc:
            full_text += page.get_text() + "\n"
    os.unlink(temp_path)
    if len(full_text.strip()) < 500:
        logger.info(f"Low text detected, using OCR for {doc_id}")
        # OCR fallback: Upload to temp storage, annotate
        temp_blob_name = f"temp/{doc_id}.pdf"
        input_blob = storage_client.bucket(bucket_name).blob(temp_blob_name)
        input_blob.upload_from_string(content, content_type="application/pdf")
        output_uri = f"gs://{bucket_name}/temp/{doc_id}_ocr/"
        request = vision.BatchAnnotateFilesRequest(
            requests=[
                vision.AnnotateFileRequest(
                    input_config=vision.InputConfig(
                        gcs_source=vision.GcsSource(uri=f"gs://{bucket_name}/{temp_blob_name}"),
                        mime_type="application/pdf"
                    ),
                    features=[vision.Feature(type_=vision.Feature.Type.DOCUMENT_TEXT_DETECTION)],
                    output_config=vision.OutputConfig(
                        gcs_destination=vision.GcsDestination(uri=output_uri),
                        batch_size=1
                    )
                )
            ]
        )
        operation = vision_client.batch_annotate_files(request)
        # Wait for operation (simplified; in prod, poll)
        import time
        while not operation.done():
            time.sleep(5)
            operation = vision_client.transport.operations_client.get_operation(operation.name)
        # Download output JSON
        output_blob_name = f"temp/{doc_id}_ocr/output-1-to-1.json"
        output_blob = storage_client.bucket(bucket_name).blob(output_blob_name)
        if output_blob.exists():
            output_content = output_blob.download_as_text()
            data = json.loads(output_content)
            responses = data.get('responses', [])
            full_text = ""
            for resp in responses:
                annotation = resp.get('fullTextAnnotation', {})
                full_text += annotation.get('text', '') + "\n"
            if not full_text.strip():
                logger.warning(f"No text extracted via OCR for {doc_id}")
                full_text = "No text could be extracted from the document."
        else:
            logger.error(f"OCR output not found for {doc_id}")
            full_text = "OCR processing failed."
        # Cleanup blobs
        input_blob.delete()
        if output_blob.exists():
            output_blob.delete()
        # Cleanup output dir if needed (recursive delete via gsutil or list/delete)
    else:
        logger.info(f"Text extracted via PyMuPDF for {doc_id}")
    # Upload PDF to Storage
    blob = storage_client.bucket(bucket_name).blob(f"{user_id}/{doc_id}.pdf")
    blob.upload_from_string(content, content_type="application/pdf")
    # AI analysis
    try:
        facts = await generate_content(PROMPT1.format(full_text=full_text))
        analysis = await generate_content(PROMPT2.format(facts=facts, full_text=full_text))
        final_brief = await generate_content(PROMPT3.format(facts=facts, analysis=analysis, full_text=full_text))
        doc_ref.update({
            "status": "completed",
            "brief": final_brief,
            "facts": facts,
            "analysis": analysis
        })
    except Exception as e:
        doc_ref.update({"status": "failed", "error": str(e)})
        raise HTTPException(500, "Processing failed")
    return {"docId": doc_id}

@app.get("/v1/export/{doc_id}")
async def export_brief(doc_id: str, user_id: str = Depends(get_current_user)):
    logger.info(f"Exporting brief {doc_id} for user {user_id}")
    doc_ref = db.collection("users").document(user_id).collection("documents").document(doc_id)
    doc = doc_ref.get()
    if not doc.exists:
        logger.warning(f"Document {doc_id} not found for user {user_id}")
        raise HTTPException(404, "Document not found")
    data = doc.to_dict()
    if data["status"] != "completed":
        raise HTTPException(400, "Document not ready")
    brief = data["brief"]
    # Parse sections with regex for robust extraction
    headings = ["Facts", "Procedural History", "Issues", "Holding", "Reasoning", "Conclusion"]
    pattern = r'(Facts|Procedural History|Issues|Holding|Reasoning|Conclusion):\s*(.*?)(?=\n[A-Z][a-z]+:|\Z)'
    matches = re.findall(pattern, brief, re.DOTALL | re.IGNORECASE)
    sections = {match[0].title(): match[1].strip() for match in matches}
    # Fill missing sections with empty
    for heading in headings:
        if heading not in sections:
            sections[heading] = ""
            logger.warning(f"Missing section '{heading}' in brief {doc_id}")
    # Generate docx
    docx_doc = Document()
    for heading, content in sections.items():
        heading_p = docx_doc.add_paragraph()
        heading_run = heading_p.add_run(heading)
        heading_run.bold = True
        heading_run.font.size = Inches(0.14)
        content_p = docx_doc.add_paragraph(content)
    stream = BytesIO()
    docx_doc.save(stream)
    stream.seek(0)
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=brief_{doc_id}.docx"}
    )

if __name__ == "__main__":
    import uvicorn
    reload = os.getenv("ENV", "dev") == "dev"
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=reload)